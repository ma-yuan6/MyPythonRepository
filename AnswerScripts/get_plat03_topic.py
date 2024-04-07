#!/usr/bin/env python
# coding: utf-8
# @Author  : MJX
# @Time    : 2024/1/15 9:59
# @File    : getPlat03Topic.py

from pathlib import Path
from bs4 import BeautifulSoup
from pandas import DataFrame
from docx import Document
from docx.shared import RGBColor
from tqdm import tqdm


def parse_question_el(qustion_el, id_tuple):
    """
    用于解析每个题目，获取题目信息
    :param qustion_el: 每一个题目的 div 对象
    :param id_tuple: 已经获取的题目 id 列表
    :return: 解析出来的题目信息
    """
    question_msg = []
    id_ = qustion_el.attrs['id']
    id_ = id_.replace('question', '')  # 获取题目id
    if id_ not in id_tuple:  # 判断是否已经获取过
        question_msg.append(id_)
        qustion_title = '(' + qustion_el.select_one('.typeTitle').text.split('【')[1].split('】')[0] + ')'
        qustion = qustion_title + qustion_el.select_one('.type').text
        question_msg.append(qustion.replace('题)', ')'))  # 题目
        anwser = qustion_el.select('H4')[1].text.split('：')[1]  # 答案
        question_msg.append(anwser)
        options_el = qustion_el.select('.optionitem')  # 选项
        for option in options_el:
            question_msg.append(option.text.replace('\xa0\xa0', ''))
    return question_msg


file_list = Path(__file__).parent.joinpath('source').iterdir()  # 源数据文件
id_tuple = tuple()  # 存储题目 id 列表
all_question = []  # 存储所有题目
tqdm_bar = tqdm(list(file_list))
for file in tqdm_bar:
    tqdm_bar.set_description('Parsing ' + file.name)
    with open(file, 'r', encoding='utf-8') as f:
        html = f.read()
    soup = BeautifulSoup(html, 'lxml')
    qustion_el_list = soup.select('#allQuestion .box')  # 题目 div
    for qustion_el in qustion_el_list:
        question_msg = parse_question_el(qustion_el, id_tuple=id_tuple)  # 解析
        if question_msg:
            all_question.append(tuple(question_msg))
            id_tuple = id_tuple + (question_msg[0],)
all_question.sort(key=lambda x: int(x[0]))  # 按照 id 排序
print('获取题目数：', len(all_question))
# print(all_question)
# 存储数据
df = DataFrame(all_question, columns=['小节标题', '题目', '答案', '选项一', '选项二', '选项三', '选项四', '选项五'])
df['小节标题'] = '训练题' + df['小节标题']
df.to_excel(Path(__file__).parent.joinpath('platform3').joinpath('platform3_all.xlsx'), index=False)

if not Path(__file__).parent.joinpath('platform3').exists():
    Path(__file__).parent.joinpath('platform3').mkdir()
with open(Path(__file__).parent.joinpath('platform3').joinpath('platform3_all.txt'),
          'w', encoding='utf-8') as f:
    for question in all_question:
        f.write('    '.join(question))
        f.write('\n')

doc = Document()
doc.add_heading('平台三题目汇总', 0)  # 大标题
idx = 1
for question in all_question:

    topic = doc.add_paragraph(str(idx) + '、' + question[1])
    ans = topic.add_run(f'({question[2]})')  # 写入答案
    ans.bold = True  # 答案加粗
    ans.font.color.rgb = RGBColor(169, 60, 4)  # 答案颜色
    for option in question[3:]:
        doc.add_paragraph(option)  # 选项
    idx += 1
doc.save(Path(__file__).parent.joinpath('platform3').joinpath('platform3_all.docx'))
