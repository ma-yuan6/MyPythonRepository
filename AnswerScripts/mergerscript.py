#!/usr/bin/env python
# coding: utf-8
# @Author  : MJX
# @Time    : 2023/12/16 12:34
# @File    : mergerScript.py

import os
from tqdm import tqdm
from docx import Document
from docx.shared import RGBColor
import pandas as pd
from settings import DIR

HEADER = ''
if DIR == 'platform1':
    HEADER = '平台一题目汇总'
elif DIR == 'platform2':
    HEADER = '平台二题目汇总'

topic_files = os.listdir(DIR)
topic_files = [file for file in topic_files if file[0] == '0']

# 汇总所有题目
f_all = open(f'./{DIR}/{DIR}_all.txt', 'w', encoding='utf-8')
for file in tqdm(topic_files):
    with open(f'{DIR}/{file}', 'r', encoding='utf-8') as f:
        for line in f:
            line = '(' + line.split('、(')[1]
            f_all.write(file.split('-')[1].rstrip('.txt'))
            f_all.write('    ')
            f_all.write(line)
        f_all.write('\n')
f_all.close()

# 转换为excel
df = pd.read_csv(f'./{DIR}/{DIR}_all.txt', sep='    ', header=None,
                 encoding='utf-8', skip_blank_lines=True, engine='python',
                 names=['小节标题', '题目', '答案', '选项一', '选项二', '选项三', '选项四'],
                 )
df.to_excel(f'./{DIR}/{DIR}_all.xlsx', index=False)

# 转换为 word
f_all = open(f'./{DIR}/{DIR}_all.txt', 'r', encoding='utf-8')
doc = Document()
doc.add_heading(HEADER, 0)  # 大标题
title = ''
idx = 1
for line in f_all:
    if line == '\n':  # 遇到空格行就换标题
        title = ''
        idx = 1
        continue
    topic_list = line.split('    ')

    if not title:
        title = topic_list[0]  # 二级标题
        doc.add_heading(title, level=1)
    topic = doc.add_paragraph(str(idx) + '、' + topic_list[1])  # 写入题目
    ans = topic.add_run(f'({topic_list[2]})')  # 写入答案
    ans.bold = True  # 答案加粗
    ans.font.color.rgb = RGBColor(169, 60, 4)  # 答案颜色

    for option in topic_list[3:]:  # 写入选项
        doc.add_paragraph(option)
    idx += 1
doc.save(f'./{DIR}/{DIR}_all.docx')
